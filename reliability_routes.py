"""
reliability_routes.py — Flask Blueprint: Reliability Performance & Risk Agent
Endpoints: upload laporan bulanan, run agent, history output.
"""

import os
import io
from datetime import datetime
from flask import Blueprint, request, jsonify, send_from_directory

from docx import Document
from reliability_agent import run_reliability_agent
from reliability_data import save_laporan_bulanan, ensure_reliability_schema, get_source_rows, get_dashboard_data
from db import (
    save_reliability_output,
    fetch_reliability_outputs,
    fetch_reliability_output_detail,
)

reliability_bp = Blueprint("reliability", __name__)

STATIC_DIR = os.path.join(os.path.dirname(__file__), "static")

# Jalankan schema patch saat blueprint dimuat
try:
    ensure_reliability_schema()
except Exception as e:
    print(f"[Reliability] Schema patch warning: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# SERVE UI
# ─────────────────────────────────────────────────────────────────────────────

@reliability_bp.route("/reliability")
def reliability_ui():
    resp = send_from_directory(STATIC_DIR, "reliability.html")
    # Cegah browser cache UI lama agar update deploy langsung kepakai
    resp.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


# ─────────────────────────────────────────────────────────────────────────────
# UPLOAD LAPORAN BULANAN
# ─────────────────────────────────────────────────────────────────────────────

@reliability_bp.route("/reliability/upload-laporan", methods=["POST"])
def upload_laporan():
    """
    Upload file .docx laporan bulanan reliability.
    Extract teks + tabel → simpan ke tabel reports.
    """
    if "file" not in request.files:
        return jsonify({"error": "Tidak ada file yang dikirim"}), 400

    file = request.files["file"]

    if not file or file.filename == "":
        return jsonify({"error": "File tidak valid"}), 400

    if not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "Hanya file .docx yang diterima"}), 400

    try:
        doc = Document(io.BytesIO(file.read()))

        paragraphs = []

        # Extract paragraf
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract tabel
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        content = "\n".join(paragraphs)

        if not content:
            return jsonify({"error": "File kosong atau tidak bisa dibaca"}), 400

        base_name = os.path.splitext(file.filename)[0].replace("_", " ")
        title = f"{base_name} — Upload {datetime.now().strftime('%d %b %Y %H:%M')}"

        report_id = save_laporan_bulanan(title=title, content=content)

        return jsonify({
            "success":    True,
            "report_id":  report_id,
            "title":      title,
            "chars":      len(content),
            "paragraphs": len(paragraphs),
            "message":    "Laporan berhasil diupload dan siap digunakan oleh agent.",
        })

    except Exception as e:
        return jsonify({"error": f"Gagal memproses file: {str(e)}"}), 500


# ─────────────────────────────────────────────────────────────────────────────
# RUN AGENT
# ─────────────────────────────────────────────────────────────────────────────

@reliability_bp.route("/reliability/run", methods=["POST"])
def run_agent():
    """
    Jalankan Reliability Performance & Risk Agent.
    Body JSON: { "mode": "weekly" | "monthly" }
    """
    body = request.get_json(silent=True) or {}
    mode = body.get("mode", "weekly")
    ru   = body.get("ru") or None  # None = overall

    if mode not in ("weekly", "monthly"):
        return jsonify({"error": "mode harus 'weekly' atau 'monthly'"}), 400

    try:
        result = run_reliability_agent(mode=mode, ru=ru)

        now  = datetime.now()
        week = now.isocalendar()[1]
        ru_suffix = f" — {ru}" if ru else ""

        if mode == "weekly":
            title       = f"Weekly Reliability Review — W{week} {now.strftime('%b %Y')}{ru_suffix}"
            output_type = f"reliability_weekly{'_' + ru.replace(' ', '_').lower() if ru else ''}"
        else:
            title       = f"Monthly Reliability Health Review — {now.strftime('%B %Y')}{ru_suffix}"
            output_type = f"reliability_monthly{'_' + ru.replace(' ', '_').lower() if ru else ''}"

        output_id = save_reliability_output(
            output_type=output_type,
            title=title,
            content=result["content"],
            batch_ref=f"{mode}_{now.strftime('%Y%m%d_%H%M')}",
            dashboard_html=result.get("dashboard_html", ""),
        )

        return jsonify({
            "success":        True,
            "output_id":      output_id,
            "title":          title,
            "mode":           mode,
            "ru":             ru,
            "content":        result["content"],
            "dashboard_html": result.get("dashboard_html", ""),
            "dashboard_error": result.get("dashboard_error", ""),
        })

    except Exception as e:
        return jsonify({"error": f"Agent gagal dijalankan: {str(e)}"}), 500


# ─────────────────────────────────────────────────────────────────────────────
# HISTORY
# ─────────────────────────────────────────────────────────────────────────────

@reliability_bp.route("/reliability/history", methods=["GET"])
def get_history():
    mode  = request.args.get("mode", "").strip()
    limit = min(int(request.args.get("limit", 20)), 100)

    output_type = f"reliability_{mode}" if mode else None

    try:
        rows = fetch_reliability_outputs(output_type=output_type, limit=limit)
        return jsonify({"success": True, "data": [dict(r) for r in rows]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@reliability_bp.route("/reliability/history/<int:output_id>", methods=["GET"])
def get_history_detail(output_id):
    try:
        row = fetch_reliability_output_detail(output_id)
        if not row:
            return jsonify({"error": "Output tidak ditemukan"}), 404
        return jsonify({"success": True, "data": dict(row)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@reliability_bp.route("/reliability/source-data/<source_key>", methods=["GET"])
def get_source_data(source_key):
    """Kembalikan baris data mentah untuk modal 'Lihat Sumber Data'."""
    ru = request.args.get("ru", "").strip() or None
    try:
        rows, columns, title = get_source_rows(source_key, ru=ru)
        return jsonify({
            "success": True,
            "key":     source_key,
            "title":   title,
            "columns": columns,
            "rows":    rows,
            "count":   len(rows),
        })
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500



@reliability_bp.route("/reliability/dashboard-stats", methods=["GET"])
def get_dashboard_stats():
    """Kembalikan data agregasi untuk infografis dashboard (live dari DB)."""
    try:
        data = get_dashboard_data()
        return jsonify({"success": True, "data": data})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@reliability_bp.route("/reliability/last-report", methods=["GET"])
def get_last_report():
    """Ambil info laporan bulanan terakhir yang diupload."""
    try:
        from db import fetch_reports
        rows = fetch_reports(report_type="monthly_reliability", limit=1)
        if rows:
            r = dict(rows[0])
            r.pop("content", None)  # jangan kirim content penuh
            return jsonify({"success": True, "data": r})
        return jsonify({"success": True, "data": None})
    except Exception as e:
        return jsonify({"error": str(e)}), 500